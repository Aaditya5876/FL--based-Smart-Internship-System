import json
import os
from datetime import datetime


def load_rq2_results():
    agg_path = os.path.join('src', 'research', 'RQ2', 'logs', 'agg_updated.json')
    if not os.path.exists(agg_path):
        # try to locate any agg_*.json
        logs_dir = os.path.join('src', 'research', 'RQ2', 'logs')
        if os.path.isdir(logs_dir):
            cands = [os.path.join(logs_dir, f) for f in os.listdir(logs_dir) if f.startswith('agg_') and f.endswith('.json')]
            if cands:
                agg_path = max(cands, key=lambda p: os.path.getmtime(p))
    if not os.path.exists(agg_path):
        return None
    with open(agg_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def build_markdown(md_path: str, rq2):
    ts = datetime.now().strftime('%Y-%m-%d %H:%M')
    lines = []
    lines.append(f"Generated: {ts}\n")
    lines.append("# Paper Draft\n")
    lines.append("See docs/paper_draft.md for the full narrative.\n")
    if rq2:
        lines.append("\n## RQ2 Intrinsic Summary\n")
        for row in rq2.get('results', []):
            lines.append(f"- {row['version']}: F1={row['f1']:.3f}, AUC={row['auc']:.3f}, purity={row['purity']:.3f}, NMI={row['nmi']:.3f}, phrases={row.get('num_phrases')} (top_pairs={row.get('top_pairs_used')}, gold_pairs={row.get('gold_pairs')})")
    os.makedirs(os.path.dirname(md_path), exist_ok=True)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def build_docx(docx_path: str, rq2):
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:
        print('[WARN] python-docx not installed; writing markdown summary instead.')
        build_markdown(docx_path.replace('.docx', '.md'), rq2)
        return

    doc = Document()
    doc.add_heading('Privacy-Preserving Federated Internship Recommendation', 0)
    doc.add_paragraph('Heterogeneity Handling (RQ1) and Semantic Alignment (RQ2)')
    doc.add_paragraph('Author: <Your Name>')

    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'We study federated learning (FL) for internship/job recommendation under extreme heterogeneity and wording mismatches. '
        'Enhanced PFL outperforms FedAvg/FedProx/FedOpt (RQ1). A privacy-preserving contrastive alignment pipeline (RQ2) '
        'achieves strong AUC/F1 centrally and competitively in FL, sharing only weights/gradients.'
    )

    doc.add_heading('RQ2 Intrinsic Results', level=1)
    if rq2 and rq2.get('results'):
        tbl = doc.add_table(rows=1, cols=8)
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Version'
        hdr[1].text = 'F1'
        hdr[2].text = 'AUC'
        hdr[3].text = 'Purity'
        hdr[4].text = 'NMI'
        hdr[5].text = '#Phrases'
        hdr[6].text = 'TopPairs'
        hdr[7].text = 'GoldPairs'
        for row in rq2['results']:
            r = tbl.add_row().cells
            r[0].text = str(row.get('version'))
            r[1].text = f"{row.get('f1'):.3f}" if row.get('f1') is not None else ''
            r[2].text = f"{row.get('auc'):.3f}" if row.get('auc') is not None else ''
            r[3].text = f"{row.get('purity'):.3f}" if row.get('purity') is not None else ''
            r[4].text = f"{row.get('nmi'):.3f}" if row.get('nmi') is not None else ''
            r[5].text = str(row.get('num_phrases'))
            r[6].text = str(row.get('top_pairs_used'))
            r[7].text = str(row.get('gold_pairs'))
    else:
        doc.add_paragraph('No RQ2 results found. Run the evaluation scripts to generate agg_updated.json.')

    doc.add_heading('Notes', level=2)
    doc.add_paragraph('This document is auto-generated. Edit docs/paper_draft.md for narrative content and finalize figures.')

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f'[OK] Wrote {docx_path}')


def main():
    rq2 = load_rq2_results()
    # Always write a docx attempt and a markdown summary fallback
    build_docx(os.path.join('docs', 'paper_auto.docx'), rq2)
    build_markdown(os.path.join('docs', 'paper_auto.md'), rq2)


if __name__ == '__main__':
    main()

